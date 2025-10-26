"""Word 文档解析器"""

import io
import logging
from typing import Dict

import docx

from .base import BaseParser

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Word 文档解析器 (DOCX/DOC)"""

    async def parse(self, file_data: bytes, **kwargs) -> str:
        """解析 Word 文档"""
        try:
            doc = docx.Document(io.BytesIO(file_data))

            text_parts = []

            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(table_text)

            text = "\n\n".join(text_parts)
            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Failed to parse Word document: {e}")
            raise

    def _extract_table(self, table) -> str:
        """提取表格内容"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        return "\n".join(rows)

    def extract_metadata(self, file_data: bytes) -> Dict:
        """提取 Word 元数据"""
        try:
            doc = docx.Document(io.BytesIO(file_data))
            core_properties = doc.core_properties

            return {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "comments": core_properties.comments or "",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
        except Exception as e:
            logger.warning(f"Failed to extract Word metadata: {e}")
            return {}
